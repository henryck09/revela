from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Manual_operativo_Revela.docx"

NAVY = "30264D"
PURPLE = "756B91"
CREAM = "F7F3EB"
LINE = "E3D9C6"
BLUE = "2E74B5"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), bold=True)
        font(p.add_run(text[len(bold_prefix):]))
    else:
        font(p.add_run(text))
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt({1:18,2:14,3:10}[level])
    p.paragraph_format.space_after = Pt({1:10,2:7,3:5}[level])
    run = p.add_run(text)
    font(run, size={1:16,2:13,3:12}[level], color={1:BLUE,2:BLUE,3:"1F4D78"}[level], bold=True)
    return p

def key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_widths(table, [1.875, 4.625])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], "E8EEF5")
        font(cells[0].paragraphs[0].add_run(label), bold=True)
        font(cells[1].paragraphs[0].add_run(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    font(run, size=9, color=PURPLE)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

def setup(doc):
    sec = doc.sections[0]
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
        s.header_distance = Inches(.492); s.footer_distance = Inches(.492)
        header = s.header.paragraphs[0]
        header.text = "REVELA  |  Manual operativo y tecnico"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs: font(run, size=8.5, color=PURPLE, bold=True)
        add_page_number(s.footer.paragraphs[0])
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    for n in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[n].font.name = "Calibri"

def main():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document(); setup(doc)
    # Cover - editorial cover pattern with compact reference guide body tokens.
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("REVELA"), size=13, color=PURPLE, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Manual operativo\ny tecnico"), size=30, color=NAVY, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Pedidos, tarjetas QR, administracion, despliegue y mantenimiento"), size=13, color=PURPLE)
    doc.add_paragraph().paragraph_format.space_after = Pt(105)
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_widths(t, [6.5]); set_cell_shading(t.cell(0,0), CREAM)
    p = t.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Version final de trabajo | Julio 2026\nURL publica: https://revela-liart.vercel.app"), size=11, color=NAVY, bold=True)
    doc.add_page_break()

    heading(doc, "Indice", 1)
    for item in ["1. Resumen y objetivo", "2. Arquitectura y componentes", "3. Flujo operativo de un pedido", "4. Tarjeta QR y reenvio manual", "5. Panel de administracion", "6. Datos y almacenamiento", "7. Configuracion y despliegue", "8. Mantenimiento y solucion de problemas", "9. Seguridad y limites", "10. Lista de verificacion final"]:
        bullet(doc, item)

    heading(doc, "1. Resumen y objetivo", 1)
    para(doc, "Revela es una aplicacion web para crear capsulas digitales personalizadas. Cada pedido contiene texto, fotos, musica, video, fecha y estilo visual. Tras aprobar el pedido, se genera una tarjeta cuadrada con un codigo QR; al escanearlo, la persona destinataria abre la capsula.")
    key_value_table(doc, [("Sitio publico", "https://revela-liart.vercel.app"), ("Creador", "Pagina inicial /"), ("Capsula", "Ruta /m/CODIGO"), ("Administrador", "Ruta /admin/login y /admin"), ("Base de datos y archivos", "Supabase"), ("Correo", "Resend, con reenvio manual temporal")])

    heading(doc, "2. Arquitectura y componentes", 1)
    para(doc, "La aplicacion usa React y Vite para la interfaz, Supabase para autenticacion, tabla de pedidos y Storage, Resend para el correo y Vercel para publicar la web.")
    heading(doc, "Rutas principales", 2)
    key_value_table(doc, [("/", "Formulario de creacion y vista previa de la capsula."), ("/m/:code", "Pagina publica que carga un pedido aprobado y revela su contenido."), ("/admin/login", "Inicio de sesion del administrador de Supabase."), ("/admin", "Panel de control de pedidos, envios y eliminacion.")])
    heading(doc, "Archivos principales", 2)
    key_value_table(doc, [("src/pages/CreatorPage.jsx", "Formulario de pedido, carga de medios y vista previa."), ("src/pages/RevealPage.jsx", "Carga publica y animacion de revelado."), ("src/components/CapsuleStory.jsx", "Contenido de la capsula y controles de musica."), ("src/pages/AdminDashboard.jsx", "Control, busqueda, filtros, periodos y acciones administrativas."), ("src/services/orders.js", "Operaciones de pedidos y llamadas a funciones."), ("supabase/functions", "Funciones de correo y eliminacion segura."), ("supabase/schema.sql", "Esquema inicial de base de datos y Storage.")])

    heading(doc, "3. Flujo operativo de un pedido", 1)
    for step in ["La persona completa el formulario y adjunta los medios permitidos.", "La aplicacion sube fotos, audio, video y fondos al bucket capsule-media y guarda el pedido como PENDIENTE.", "El cliente envia la captura de pago por WhatsApp.", "El administrador abre /admin, revisa el pedido y pulsa Aprobar y enviar.", "El pedido queda APROBADO. Resend envia la tarjeta QR al correo personal del administrador.", "El administrador reenvia manualmente la tarjeta al correo que dejo el cliente, o la imprime para entregarla.", "La persona destinataria escanea el QR y ve la capsula en /m/CODIGO."]:
        bullet(doc, step)

    heading(doc, "4. Tarjeta QR y reenvio manual", 1)
    para(doc, "Mientras no exista un dominio verificado en Resend, la cuenta de prueba solo puede enviar a la direccion asociada a Resend. Por eso la funcion envia la tarjeta al ADMIN_EMAIL; el correo que el cliente escribe se conserva como referencia del pedido.")
    heading(doc, "Caracteristicas de la tarjeta", 2)
    for item in ["Formato cuadrado de 480 x 480 px dentro del correo.", "Mensaje central: Alguien preparo algo bonito para ti.", "QR grande incrustado como adjunto CID; no depende de descargar imagenes externas al imprimir.", "No muestra la URL ni el texto lista para imprimir dentro de la tarjeta.", "El asunto del correo contiene codigo y nombre del pedido para identificar a quien reenviar."]:
        bullet(doc, item)
    heading(doc, "Procedimiento de reenvio", 2)
    for step in ["Aprueba el pedido en el panel.", "Abre el correo recibido en ADMIN_EMAIL.", "Comprueba el codigo de pedido y el nombre del cliente en el asunto.", "Reenvia el correo al cliente o imprime solo la tarjeta cuadrada.", "Prueba el QR con otro telefono antes de la entrega."]:
        bullet(doc, step)
    para(doc, "Nota: abrir el enlace alternativo desde un telefono muestra la capsula directamente. Eso es correcto; la experiencia de regalo se logra cuando otra persona escanea el QR impreso.")

    heading(doc, "5. Panel de administracion", 1)
    para(doc, "El panel requiere iniciar sesion con un usuario creado en Supabase Authentication. Sus funciones actuales son:")
    for item in ["Indicadores de total de pedidos, pendientes, aprobados e ingresos aprobados.", "Busqueda por codigo, nombre, correo o WhatsApp.", "Filtro por pendiente, aprobado o rechazado.", "Seccion de pendientes y un historial agrupado por esta semana, semana pasada y meses anteriores.", "Ficha con fecha, precio, ocasion, archivos asociados, correo, WhatsApp y enlace de capsula.", "Aprobacion, rechazo, reenvio de tarjeta, copia de correo y eliminacion.", "Eliminacion con confirmacion: borra el pedido y sus fotos, audio, video y fondo asociados."]:
        bullet(doc, item)
    para(doc, "La eliminacion es irreversible. Usala para pruebas, duplicados o pedidos que deban retirarse. Un pedido eliminado deja de abrirse desde su QR.", bold_prefix="La eliminacion es irreversible.")

    heading(doc, "6. Datos y almacenamiento", 1)
    key_value_table(doc, [("Tabla orders", "Datos del pedido, estado de pago, textos, URLs de medios y fecha de creacion."), ("Bucket capsule-media", "Archivos publicos en carpetas photos/, videos/, songs/ y backgrounds/."), ("Fotos", "Se guardan como JSON en orders.photos y como archivo en Storage."), ("Seguridad publica", "Solo pedidos APROBADOS pueden leerse publicamente desde /m/:code."), ("Limpieza", "La eliminacion desde el panel borra registro y archivos del pedido.")])
    heading(doc, "Limpieza de fotos antiguas de prueba", 2)
    para(doc, "Antes de ejecutar una limpieza masiva, revisa que ningun QR antiguo deba conservar sus fotos. En Supabase SQL Editor puedes vaciar las referencias y borrar solo la carpeta photos/ con:")
    code = doc.add_paragraph(); code.paragraph_format.space_after = Pt(8); code.paragraph_format.left_indent = Inches(.2)
    font(code.add_run("update public.orders set photos = '[]'::jsonb where jsonb_array_length(photos) > 0;\n\ndelete from storage.objects where bucket_id = 'capsule-media' and name like 'photos/%';"), name="Consolas", size=9, color=NAVY)

    heading(doc, "7. Configuracion y despliegue", 1)
    heading(doc, "Variables publicas de Vercel", 2)
    key_value_table(doc, [("VITE_SUPABASE_URL", "URL del proyecto Supabase."), ("VITE_SUPABASE_ANON_KEY", "Clave anon publica de Supabase."), ("VITE_WHATSAPP_NUMBER", "Numero del negocio en formato internacional, sin +."), ("VITE_SITE_URL", "https://revela-liart.vercel.app")])
    heading(doc, "Secretos de Edge Functions", 2)
    key_value_table(doc, [("RESEND_API_KEY", "Clave privada de Resend. No subir a GitHub."), ("SITE_URL", "https://revela-liart.vercel.app"), ("FROM_EMAIL", "Revela <onboarding@resend.dev> durante pruebas."), ("ADMIN_EMAIL", "henry-campa_10@hotmail.com para recibir las tarjetas.")])
    heading(doc, "Comandos de despliegue", 2)
    for command in ["git add .", "git commit -m \"Actualiza version final Revela\"", "git push", "supabase secrets set ADMIN_EMAIL=\"henry-campa_10@hotmail.com\"", "supabase functions deploy send-capsule-email", "supabase functions deploy delete-order"]:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.2); p.paragraph_format.space_after = Pt(3); font(p.add_run(command), name="Consolas", size=9, color=NAVY)
    para(doc, "Vercel publica automaticamente la interfaz tras el git push si el repositorio esta conectado. Las Edge Functions de Supabase siempre se despliegan con sus comandos propios.")

    heading(doc, "8. Mantenimiento y solucion de problemas", 1)
    key_value_table(doc, [("No llega el correo", "Revisar logs de send-capsule-email en Supabase. Confirmar RESEND_API_KEY, FROM_EMAIL, SITE_URL y ADMIN_EMAIL."), ("Llega solo al correo personal", "Es el comportamiento esperado de onboarding@resend.dev. Para clientes reales se necesita dominio verificado."), ("QR no se ve", "La tarjeta usa imagen CID adjunta. Revisar que el cliente de correo permita adjuntos y probar la imagen PNG adjunta."), ("QR abre error", "Confirmar SITE_URL y que el pedido este APROBADO."), ("No se borra pedido", "Desplegar delete-order y confirmar sesion de administrador valida."), ("Almacenamiento lleno", "Eliminar pedidos de prueba desde /admin. Esto borra los archivos asociados."), ("Musica no inicia", "La musica de YouTube se reproduce como audio con una barra. Requiere pulsar reproducir; algunas pistas pueden restringir embeds.")])

    heading(doc, "9. Seguridad y limites", 1)
    for item in ["Nunca subir .env, claves de Resend ni claves service-role a GitHub.", "Cambiar la contrasena del administrador si se comparte o se sospecha acceso no autorizado.", "El bucket es publico porque las capsulas usan sus archivos directamente; no almacenar documentos privados ni informacion sensible.", "Eliminar pedidos desde el panel es irreversible y deja inutilizable su QR.", "El plan de Resend sin dominio es solo de prueba. Antes de automatizar envios al cliente, verificar un dominio propio y cambiar FROM_EMAIL.", "Mantener limites de archivo en el creador para no agotar Storage."]:
        bullet(doc, item)

    heading(doc, "10. Lista de verificacion final", 1)
    for item in ["Las variables de Vercel estan configuradas.", "Los cuatro secretos de Supabase estan configurados, incluido ADMIN_EMAIL.", "send-capsule-email y delete-order estan desplegadas.", "Se creo un pedido de prueba con fotos y musica.", "El administrador aprobo el pedido y recibio la tarjeta en Hotmail.", "El QR impreso se probo desde otro telefono.", "La capsula muestra fotos y musica sin video de YouTube.", "Se probo reenviar tarjeta y eliminar un pedido de prueba.", "Los cambios estan en GitHub y Vercel muestra el despliegue correcto."]:
        bullet(doc, "[ ] " + item)

    doc.save(OUT)
    print(OUT)

if __name__ == "__main__": main()
